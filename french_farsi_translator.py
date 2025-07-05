#!/usr/bin/env python3
"""
French to Farsi Translation Tool
Created for: ehsan42324232/french-farsi-translation

This script processes French text and creates translated Word documents
with proper RTL (Right-to-Left) formatting for Farsi text.

Requirements:
- python-docx
- googletrans (for translation)
- requests

Install with: pip install python-docx googletrans==4.0.0rc1 requests
"""

import os
import re
import time
import json
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("Please install python-docx: pip install python-docx")

try:
    from googletrans import Translator
except ImportError:
    print("Please install googletrans: pip install googletrans==4.0.0rc1")

class FrenchFarsiTranslator:
    def __init__(self, output_dir="translations"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.translator = Translator()
        
        # Common French to Farsi translations for consistency
        self.custom_translations = {
            "bonjour": "سلام",
            "bonsoir": "عصر بخیر", 
            "bonne nuit": "شب بخیر",
            "au revoir": "خداحافظ",
            "merci": "متشکرم",
            "merci beaucoup": "بسیار متشکرم",
            "s'il vous plaît": "لطفاً",
            "s'il te plaît": "لطفاً",
            "excusez-moi": "ببخشید",
            "pardon": "ببخشید",
            "comment allez-vous": "حالتان چطور است",
            "comment ça va": "چطوری",
            "ça va": "خوبم",
            "très bien": "خیلی خوب",
            "bien": "خوب",
            "oui": "بله",
            "non": "نه",
            "peut-être": "شاید",
            "aujourd'hui": "امروز",
            "hier": "دیروز", 
            "demain": "فردا",
            "maintenant": "الان",
            "toujours": "همیشه",
            "jamais": "هرگز",
            "beaucoup": "خیلی",
            "un peu": "کمی",
            "grand": "بزرگ",
            "petit": "کوچک",
            "nouveau": "جدید",
            "vieux": "قدیمی",
            "jeune": "جوان",
            "important": "مهم",
            "facile": "آسان",
            "difficile": "سخت",
            "possible": "ممکن",
            "impossible": "غیرممکن"
        }
    
    def split_text_into_chunks(self, text, max_chars=4000):
        """Split text into chunks for translation, preserving sentence boundaries."""
        sentences = re.split(r'[.!?]+', text)
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            # If adding this sentence would exceed max_chars, start a new chunk
            if len(current_chunk + sentence) > max_chars and current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = sentence
            else:
                current_chunk += " " + sentence if current_chunk else sentence
        
        # Don't forget the last chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def translate_chunk(self, text_chunk, delay=1):
        """Translate a chunk of text from French to Farsi."""
        try:
            # Apply custom translations first
            modified_text = text_chunk.lower()
            for french, farsi in self.custom_translations.items():
                pattern = r'\b' + re.escape(french) + r'\b'
                modified_text = re.sub(pattern, farsi, modified_text, flags=re.IGNORECASE)
            
            # Use Google Translate for the rest
            time.sleep(delay)  # Rate limiting
            result = self.translator.translate(text_chunk, src='fr', dest='fa')
            
            return result.text if result else text_chunk
            
        except Exception as e:
            print(f"Translation error: {e}")
            # Fallback to basic translation
            return self.basic_translate(text_chunk)
    
    def basic_translate(self, text):
        """Fallback basic translation using dictionary."""
        translated = text.lower()
        for french, farsi in self.custom_translations.items():
            pattern = r'\b' + re.escape(french) + r'\b'
            translated = re.sub(pattern, farsi, translated, flags=re.IGNORECASE)
        return translated
    
    def create_word_document(self, original_text, translated_chunks, include_original=True):
        """Create a Word document with proper RTL formatting for Farsi."""
        doc = Document()
        
        # Set up document styles
        self.setup_document_styles(doc)
        
        # Document title and metadata
        title = doc.add_heading('ترجمه از فرانسه به فارسی', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.set_rtl_paragraph(title)
        
        # Add metadata
        metadata_para = doc.add_paragraph()
        metadata_para.add_run(f'تاریخ ترجمه: {datetime.now().strftime("%Y/%m/%d")}').bold = True
        metadata_para.add_run(f'\nتعداد بخش‌ها: {len(translated_chunks)}')
        metadata_para.add_run(f'\nطول متن اصلی: {len(original_text):,} کاراکتر')
        metadata_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self.set_rtl_paragraph(metadata_para)
        
        # Separator
        doc.add_paragraph('═' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Process each translated chunk
        for i, chunk_data in enumerate(translated_chunks, 1):
            # Section header
            section_header = doc.add_heading(f'بخش {i}', level=2)
            section_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            self.set_rtl_paragraph(section_header)
            
            if include_original:
                # Original French text
                original_label = doc.add_paragraph('متن اصلی (فرانسه):')
                original_label.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                original_label.runs[0].bold = True
                self.set_rtl_paragraph(original_label)
                
                original_para = doc.add_paragraph(chunk_data['original'])
                original_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                original_para.style.font.name = 'Times New Roman'
                original_para.style.font.size = Pt(11)
                
                # Add some space
                doc.add_paragraph('')
            
            # Translated Farsi text
            translation_label = doc.add_paragraph('ترجمه (فارسی):')
            translation_label.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            translation_label.runs[0].bold = True
            self.set_rtl_paragraph(translation_label)
            
            farsi_para = doc.add_paragraph(chunk_data['translated'])
            farsi_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            self.set_rtl_paragraph(farsi_para)
            farsi_para.style.font.name = 'B Nazanin'
            farsi_para.style.font.size = Pt(12)
            
            # Add separator between sections
            if i < len(translated_chunks):
                doc.add_paragraph('─' * 30).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return doc
    
    def setup_document_styles(self, doc):
        """Set up document styles for better formatting."""
        try:
            # Create or modify styles for RTL text
            styles = doc.styles
            
            # Farsi paragraph style
            if 'Farsi' not in [style.name for style in styles]:
                farsi_style = styles.add_style('Farsi', WD_STYLE_TYPE.PARAGRAPH)
                farsi_style.font.name = 'B Nazanin'
                farsi_style.font.size = Pt(12)
                farsi_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
        except Exception as e:
            print(f"Style setup warning: {e}")
    
    def set_rtl_paragraph(self, paragraph):
        """Set Right-to-Left direction for Farsi text."""
        try:
            pPr = paragraph._element.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            bidi.set(qn('w:val'), '1')
            pPr.append(bidi)
        except Exception as e:
            print(f"RTL setup warning: {e}")
    
    def translate_text(self, input_text, include_original=True, chunk_size=3000):
        """Main translation function."""
        print("🚀 Starting French to Farsi translation...")
        print(f"📄 Text length: {len(input_text):,} characters")
        
        # Split text into manageable chunks
        print("✂️ Splitting text into chunks...")
        chunks = self.split_text_into_chunks(input_text, chunk_size)
        print(f"📦 Created {len(chunks)} chunks")
        
        # Translate each chunk
        translated_chunks = []
        for i, chunk in enumerate(chunks, 1):
            print(f"🔄 Translating chunk {i}/{len(chunks)}...")
            translated = self.translate_chunk(chunk)
            translated_chunks.append({
                'original': chunk,
                'translated': translated,
                'chunk_number': i
            })
            
            # Progress indicator
            progress = (i / len(chunks)) * 100
            print(f"   Progress: {progress:.1f}%")
        
        # Create Word document
        print("📝 Creating Word document...")
        doc = self.create_word_document(input_text, translated_chunks, include_original)
        
        # Save document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = self.output_dir / f"french_farsi_translation_{timestamp}.docx"
        doc.save(output_file)
        
        print(f"✅ Translation completed!")
        print(f"📁 Saved to: {output_file}")
        
        return output_file, translated_chunks

def main():
    """Example usage of the translator."""
    # Sample French text for testing
    sample_text = """
    Bonjour tout le monde. Comment allez-vous aujourd'hui? 
    J'espère que vous passez une excellente journée. 
    La traduction automatique est un domaine fascinant de l'intelligence artificielle.
    Elle permet aux gens de différentes cultures de communiquer plus facilement.
    
    C'est un exemple de texte français qui sera traduit en farsi.
    Le processus de traduction préserve la structure du document original.
    Merci beaucoup pour votre attention.
    """
    
    print("🌟 French to Farsi Translator")
    print("=" * 50)
    
    # Initialize translator
    translator = FrenchFarsiTranslator()
    
    # Option 1: Translate sample text
    print("\n📝 Translating sample text...")
    output_file, chunks = translator.translate_text(sample_text)
    
    print(f"\n📋 Translation Summary:")
    print(f"   Original length: {len(sample_text)} characters")
    print(f"   Number of chunks: {len(chunks)}")
    print(f"   Output file: {output_file}")
    
    # Option 2: Translate from file
    print("\n📁 To translate from a file, use:")
    print("   translator = FrenchFarsiTranslator()")
    print("   with open('your_french_file.txt', 'r', encoding='utf-8') as f:")
    print("       text = f.read()")
    print("   translator.translate_text(text)")

if __name__ == "__main__":
    main()
